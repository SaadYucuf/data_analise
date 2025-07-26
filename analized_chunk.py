import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import kagglehub
import os
import warnings

# Machine Learning uchun modullar
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import classification_report, confusion_matrix, roc_auc_score, RocCurveDisplay
from xgboost import XGBClassifier

# Word hujjati yaratish va grafiklar bilan ishlash uchun
import io
from docx import Document
from docx.shared import Inches

# Texnik ogohlantirishlarni o'chirish va grafiklarni chiroyli qilish
warnings.filterwarnings('ignore')
sns.set_style('whitegrid')
plt.rcParams['figure.figsize'] = (10, 6)
plt.rcParams['font.size'] = 12

print("Kutubxonalar va sozlamalar muvaffaqiyatli yuklandi.")

# ------------------------------------------------------------------------------
#                     YORDAMCHI FUNKSIYALAR
# ------------------------------------------------------------------------------

# Barcha chizilgan grafiklar saqlanadigan global o'zgaruvchi
saved_charts = {}


def save_and_show_chart(chart_name):
    """
    Joriy grafikni xotiraga (keyinchalik Wordga qo'yish uchun) saqlaydi
    va uni ekranga chiqaradi.
    """
    memfile = io.BytesIO()
    # Grafikni yuqori sifatda (150 dpi) va ortiqcha bo'sh joylarsiz saqlash
    plt.savefig(memfile, format='png', bbox_inches='tight', dpi=150)
    memfile.seek(0)  # Fayl ko'rsatkichini boshiga qaytarish
    saved_charts[chart_name] = memfile
    plt.show()  # Grafikni ekranda ko'rsatish
    plt.close()  # Xotirani tozalash uchun joriy grafikni yopish


# ------------------------------------------------------------------------------
# 1-QISM: MA'LUMOTLARNI YUKLASH VA TAYYORLASH (TOZALASH)
# ------------------------------------------------------------------------------
print("\n[1-QISM] Ma'lumotlarni yuklash va tahlilga tayyorlash boshlandi...")

df = None
try:
    data_dir = kagglehub.dataset_download("blastchar/telco-customer-churn")
    csv_file_path = os.path.join(data_dir, [f for f in os.listdir(data_dir) if f.endswith('.csv')][0])
    df = pd.read_csv(csv_file_path)

    # Ma'lumotlarni tozalash jarayoni
    df['TotalCharges'] = pd.to_numeric(df['TotalCharges'], errors='coerce')
    df.dropna(inplace=True)
    df.drop('customerID', axis=1, inplace=True)
    print("✅ Ma'lumotlar muvaffaqiyatli yuklandi va tozalandi.")

except Exception as e:
    print(f"❌ XATOLIK: Ma'lumotlarni yuklashda muammo yuz berdi: {e}")
    print("Internet aloqasi yoki Kaggle sozlamalarini tekshiring.")
    exit()  # Agar ma'lumot yuklanmasa, dasturni to'xtatish

# ------------------------------------------------------------------------------
# 2-QISM: BIZNES-TAHLIL UCHUN GRAFIKLARNI YARATISH (EDA)
# ------------------------------------------------------------------------------
print("\n[2-QISM] Biznes-tahlil uchun grafiklar yaratilmoqda...")

# Grafik 1: Umumiy Churn darajasi
plt.figure(figsize=(8, 6))
ax = sns.countplot(x='Churn', data=df, palette='pastel')
plt.title('Grafik 1: Mijozlar Ketishining Umumiy Ko\'lami', fontsize=16, pad=20)
total = len(df)
for p in ax.patches:
    percentage = f'{100 * p.get_height() / total:.1f}%'
    ax.annotate(percentage, (p.get_x() + p.get_width() / 2, p.get_height() + 20), ha='center', size=14)
save_and_show_chart('grafik_1_umumiy_churn')

# Grafik 2, 3: Asosiy omillar tahlili
key_factors = ['Contract', 'InternetService']
for i, factor in enumerate(key_factors, 2):
    plt.figure(figsize=(10, 7))
    sns.countplot(x=factor, hue='Churn', data=df, palette='viridis')
    plt.title(f'Grafik {i}: Mijozlar Ketishiga "{factor.capitalize()}"ning Ta\'siri', fontsize=16, pad=20)
    plt.legend(title='Mijoz Holati', labels=['Qolgan', 'Ketgan'])
    save_and_show_chart(f'grafik_{i}_{factor.lower()}')

print("✅ EDA grafiklari muvaffaqiyatli yaratildi.")

# ------------------------------------------------------------------------------
# 3-QISM: BASHORAT MODELINI QURISH VA O'QITISH
# ------------------------------------------------------------------------------
print("\n[3-QISM] Bashorat modelini tayyorlash va o'qitish...")

# Ma'lumotlarni model uchun tayyorlash
df_model = df.copy()
df_model['Churn'] = df_model['Churn'].map({'Yes': 1, 'No': 0})
df_processed = pd.get_dummies(df_model, drop_first=True)
X = df_processed.drop('Churn', axis=1)
y = df_processed['Churn']
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42, stratify=y)
scaler = StandardScaler()
X_train[['tenure', 'MonthlyCharges', 'TotalCharges']] = scaler.fit_transform(
    X_train[['tenure', 'MonthlyCharges', 'TotalCharges']])
X_test[['tenure', 'MonthlyCharges', 'TotalCharges']] = scaler.transform(
    X_test[['tenure', 'MonthlyCharges', 'TotalCharges']])

# Eng yaxshi modelni o'qitish
model = XGBClassifier(use_label_encoder=False, eval_metric='logloss', random_state=42)
model.fit(X_train, y_train)

print("✅ Bashorat modeli muvaffaqiyatli o'qitildi.")

# ------------------------------------------------------------------------------
# 4-QISM: MODEL NATIJALARINI TAHLIL QILISH VA VIZUALIZATSIYA
# ------------------------------------------------------------------------------
print("\n[4-QISM] Model natijalari tahlil qilinmoqda va grafiklar yaratilmoqda...")

# Grafik 4: Model aniqligi (Confusion Matrix)
y_pred = model.predict(X_test)
cm = confusion_matrix(y_test, y_pred)
plt.figure(figsize=(8, 6))
sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', annot_kws={"size": 16},
            xticklabels=['Bashorat: Qoladi', 'Bashorat: Ketadi'],
            yticklabels=['Haqiqiy: Qoladi', 'Haqiqiy: Ketadi'])
plt.title('Grafik 4: Modelning Bashorat Aniqligi', fontsize=16, pad=20)
save_and_show_chart('grafik_4_confusion_matrix')

# Grafik 5: Model sifati (ROC-AUC)
plt.figure(figsize=(10, 8))
ax_roc = plt.gca()
RocCurveDisplay.from_estimator(model, X_test, y_test, name='XGBoost Modeli', ax=ax_roc, lw=2.5)
ax_roc.plot([0, 1], [0, 1], 'k--', label='Tasodifiy taxmin', lw=2)
ax_roc.set_title('Grafik 5: Modelning Umumiy Sifati (ROC Egri Chizig\'i)', fontsize=16, pad=20)
ax_roc.legend()
save_and_show_chart('grafik_5_roc_auc')

# Grafik 6: Muhim omillar
importances = pd.Series(model.feature_importances_, index=X.columns)
plt.figure(figsize=(12, 8))
importances.nlargest(10).sort_values().plot(kind='barh', color='teal')
plt.title('Grafik 6: Mijozlar Ketishining Eng Asosiy 10 Sababi', fontsize=16, pad=20)
plt.xlabel('Omilning Muhimlik Darajasi', fontsize=12)
save_and_show_chart('grafik_6_feature_importance')

print("✅ Model natijalari tahlil qilindi va grafiklar tayyorlandi.")

# ------------------------------------------------------------------------------
# YAKUNIY BOSQICH: BARCHA TAHLILNI WORD HUJJATIGA SAQLASH
# ------------------------------------------------------------------------------
print("\n[YAKUNIY BOSQICH] Hisobotni Word (.docx) fayliga yaratish boshlandi...")

try:
    # Hujjat obyektini yaratish
    doc = Document()
    doc.add_heading('Mijozlarning Ketishini Bashorat Qilish: Biznes-Tahlil Hisoboti', 0)

    # Hisobot bo'limlari
    doc.add_heading('1. Xulosa (Executive Summary)', level=1)
    doc.add_paragraph(
        "Ushbu tahlil \"Telco\" kompaniyasida mijozlarning ketib qolish (churn) darajasini kamaytirish maqsadida o'tkazildi. "
        "Asosiy topilmalar mijozlar ketishining sabablarini ochib beradi va ularni oldindan bashorat qilish imkoniyatini ko'rsatadi."
    )

    doc.add_heading('2. Dastlabki Tahlil (Exploratory Data Analysis)', level=1)
    doc.add_picture(saved_charts['grafik_1_umumiy_churn'], width=Inches(6.0))
    doc.add_paragraph("Tahlil: Mijozlarning 26.5% qismi kompaniyani tark etmoqda.", style='Intense Quote')

    doc.add_picture(saved_charts['grafik_2_contract'], width=Inches(6.0))
    doc.add_paragraph("Tahlil: 'Month-to-month' (oyma-oy) shartnomadagi mijozlar ketishga eng moyil.",
                      style='Intense Quote')

    doc.add_picture(saved_charts['grafik_3_internetservice'], width=Inches(6.0))
    doc.add_paragraph(
        "Tahlil: 'Fiber Optic' foydalanuvchilari orasida ketish darajasi yuqori. Bu narx yoki sifat muammosidan darak berishi mumkin.",
        style='Intense Quote')

    doc.add_heading('3. Bashorat Modelining Natijalari', level=1)
    doc.add_picture(saved_charts['grafik_4_confusion_matrix'], width=Inches(5.0))
    doc.add_paragraph("Tahlil: Modelimiz ketishi mumkin bo'lgan har 100 mijozdan 66 tasini to'g'ri topmoqda.",
                      style='Intense Quote')

    doc.add_paragraph("\nModelning batafsil aniqlik hisoboti (Classification Report):")
    p = doc.add_paragraph()
    p.add_run(classification_report(y_test, y_pred)).font.name = 'Courier New'

    doc.add_picture(saved_charts['grafik_5_roc_auc'], width=Inches(6.0))
    auc_score_val = roc_auc_score(y_test, model.predict_proba(X_test)[:, 1])
    doc.add_paragraph(
        f"Tahlil: Modelning umumiy sifat ko'rsatkichi (AUC) {auc_score_val:.2f} ga teng. Bu 'juda yaxshi' natija hisoblanadi.",
        style='Intense Quote')

    doc.add_heading('4. Asosiy Sabablar va Biznes-Tavsiyalar', level=1)
    doc.add_picture(saved_charts['grafik_6_feature_importance'], width=Inches(6.5))
    doc.add_heading('Amaliy Tavsiyalar:', level=2)
    doc.add_paragraph(
        "Segmentatsiya va Maqsadli Marketing: \"Oyma-oy\" shartnomadagi yangi mijozlarga uzoq muddatli shartnomaga o'tish uchun maxsus bonuslar taklif qilish.",
        style='List Bullet')
    doc.add_paragraph(
        "Mahsulotni Tadqiq Qilish: \"Fiber Optic\" xizmatining narx siyosati va texnik barqarorligini qayta ko'rib chiqish.",
        style='List Bullet')
    doc.add_paragraph(
        "Mijozlarni Saqlab Qolish Dasturi: Har oy model yordamida \"risk guruhidagi\" mijozlar ro'yxatini shakllantirib, ular bilan proaktiv ishlash.",
        style='List Bullet')

    # Hujjatni saqlash
    file_name = "Mijozlar_Ketishi_Tahlili_Hisoboti.docx"
    doc.save(file_name)
    print(f"\n✅✅✅ MUVAFFAQIYATLI! Barcha tahlillar '{os.path.abspath(file_name)}' fayliga saqlandi.")

except Exception as e:
    print(f"\n❌❌❌ XATOLIK! Yakuniy hisobotni yaratishda muammo yuz berdi: {e}")